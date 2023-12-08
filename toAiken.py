import docx
import os
from art import *

#-----------------------------------------------------OPCIÓN 1------------------------------------------------------------------------------------

answers = []
#Añadimos la respuesta correcta a cada pregunta
def addAnswer(filename):
    
    global answers
    doc = docx.Document(filename)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold and (run.text.startswith("a)") or run.text.startswith("b)") or run.text.startswith("c)") or run.text.startswith("d)") or run.text.startswith("A)") or run.text.startswith("B)") or run.text.startswith("C)") or run.text.startswith("D)") or run.text.startswith("a.") or run.text.startswith("b.") or run.text.startswith("c.") or run.text.startswith("d.") or run.text.startswith("A.") or run.text.startswith("B.") or run.text.startswith("C.") or run.text.startswith("D.") or run.text.startswith("a ") or run.text.startswith("b ") or run.text.startswith("c ") or run.text.startswith("d ") or run.text.startswith("A ") or run.text.startswith("B ") or run.text.startswith("C ") or run.text.startswith("D ") or run.text.startswith("a-") or run.text.startswith("b-") or run.text.startswith("c-") or run.text.startswith("d-") or run.text.startswith("A-") or run.text.startswith("B-") or run.text.startswith("C-") or run.text.startswith("D-")):
                answers.append("ANSWER: " + run.text[0].upper())
            elif run.font.highlight_color != None and (run.text.startswith("a)") or run.text.startswith("b)") or run.text.startswith("c)") or run.text.startswith("d)") or run.text.startswith("A)") or run.text.startswith("B)") or run.text.startswith("C)") or run.text.startswith("D)") or run.text.startswith("a.") or run.text.startswith("b.") or run.text.startswith("c.") or run.text.startswith("d.") or run.text.startswith("A.") or run.text.startswith("B.") or run.text.startswith("C.") or run.text.startswith("D.") or run.text.startswith("a ") or run.text.startswith("b ") or run.text.startswith("c ") or run.text.startswith("d ") or run.text.startswith("A ") or run.text.startswith("B ") or run.text.startswith("C ") or run.text.startswith("D ") or run.text.startswith("a-") or run.text.startswith("b-") or run.text.startswith("c-") or run.text.startswith("d-") or run.text.startswith("A-") or run.text.startswith("B-") or run.text.startswith("C-") or run.text.startswith("D-")):
                highlight = run.text[0].upper()
                if highlight:
                    answers.append("ANSWER: " + highlight)

class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
    RED = '\033[31m'

def interface():
    print(bcolors.BOLD + bcolors.OKGREEN)
    tprint("toAiken")
    tprint("by Andressance", font="small")
    print(bcolors.BOLD + bcolors.RED + "Bienvenido al programa de transformación de documentos de word a txt en formato Aiken \n")

def interface1():
    print(bcolors.BOLD + bcolors.OKGREEN)
    tprint("toAiken")
    tprint("by Andressance", font="small")
    print(bcolors.BOLD + bcolors.RED + "Bienvenido al programa de transformación de documentos de word a txt en formato Aiken \n")
    print("Funcionamiento: \n" + bcolors.ENDC+ bcolors.OKGREEN)
    print("El programa transformará el documento word a txt, añadiendo la respuesta correcta a cada pregunta si está en negrita o subrayada en el documento word")
    print("El documento txt se guardará en la misma carpeta que el documento word")
    print("El documento txt se abrirá automáticamente al finalizar el proceso")
    print("Para comenzar, introduce el nombre del documento sin extensión")
    print("Es decir, si el documento se llama 'test.docx', introduce 'test'")
    print(bcolors.BOLD +"\n \n EJEMPLO DE FORMATEO DEL WORD: \n \n 1. ¿Cuál es la capital de España? \n a) Paris \n b) Madrid \n c) Buenos Aires \n d) Oslo \n" + bcolors.ENDC+ bcolors.OKGREEN)

#Lee el documento
def getText(filename):
    global answers
    doc = docx.Document(filename)
    fullText = []

    for para in doc.paragraphs:
        if "-" in para.text:
            para.text = para.text.replace("-", "")
        if para.text.startswith("d") or para.text.startswith("D"): 
            para.text = para.text.replace(")", ".")
            fullText.append(para.text)
            try:
                fullText.append(answers[0])
                answers.pop(0)
            except IndexError:
                pass
        else:
            para.text = para.text.replace(")", ".")
            fullText.append(para.text)
    return '\n'.join(fullText)

#Obtiene la respuesta correcta de cada pregunta

#Escribe el documento en txt
def writeText(filename, text):
    with open(filename, 'w',  encoding="utf-8") as f:
        f.write(text)

#Elimina la primera linea del word
def deleteFirstLine(filename):
    with open(filename, 'r',  encoding="utf-8") as f:
        lines = f.readlines()
    with open(filename, 'w',  encoding="utf-8") as f:
        for line in lines[1:]:
            f.write(line)
    
#Borra el numero de pregunta de cada parrafo
def deleteNumber(filename):

    with open(filename, 'r', encoding="utf-8") as f:
        lines = f.readlines()
    with open(filename, 'w',  encoding="utf-8") as f:
        for line in lines:
            if line[0].isdigit() and line[1].isdigit():
                f.write(line[3:])
            elif line[0].isdigit():
                f.write(line[3:])
            else:
                line.replace(")", ".")
                f.write(line)

#Ponemos las opciones en mayuscula
def mayusc(filename):
    
    with open(filename, 'r',  encoding="utf-8") as f:
        lines = f.readlines()
    with open(filename, 'w',  encoding="utf-8") as f:
        for line in lines:
            if line[0].startswith("a") or line[0].startswith("b") or line[0].startswith("c") or line[0].startswith("d"):
                line.replace(")", ".")
                f.write(line[:3].upper() + line[3:])
            else:
                line.replace(")", ".")
                f.write(line[0].upper() + line[1:])


#--------------------------------------------------------------------OPCION 2--------------------------------------------------------------------

abecedario = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "Ñ", "O", "P", "Q", "R", "S", "T", "U", "V", "W", "X", "Y", "Z"]

def interface2():
    print(bcolors.BOLD + bcolors.OKGREEN)
    tprint("toAiken")
    tprint("by Andressance", font="small")
    print(bcolors.BOLD + bcolors.RED + "Bienvenido al programa de creación de documentos txt en formato Aiken \n")
    print("Funcionamiento: \n" + bcolors.ENDC+ bcolors.OKGREEN)
    print("El programa transformará el documento word a txt, añadiendo la respuesta correcta a cada pregunta si está en negrita o subrayada en el documento word")
    print("Introduce el nombre del documento")
    print("Introduce el número de preguntas que quieres crear, en caso de que quieras hacer menos de las que has escrito escribe 'SALIR' y el documento se creará con las preguntas que hayas introducido")
    print("Tras introducir la pregunta, escribe el número de respuestas que quieres que tiene y posteriormente su contenido, finalmente se te pedirá la respuesta correcta")

def createQuestions():

    text = []
    numQuest = int(input("Introduce el número de preguntas a crear: "))
    for i in range(numQuest):
        pregunta = str(input(f"Introduce la pregunta {i + 1}: "))
        text.append(pregunta)
        text.append("\n")
        numAnsw = int(input("Introduce el número de respuestas de la pregunta: "))
        for j in range(numAnsw):
            res = abecedario[j] + ". "
            answ = str(input(f"Introduce la respuesta {j + 1}: " ))
            res += answ
            text.append(res)
            text.append("\n")
        correctAnsw = "ANSWER: "
        char = str(input("Introduce la respuesta correcta: "))
        char = char.strip().upper()
        correctAnsw += char
        text.append(correctAnsw)
        text.append("\n")

    return text

def createTxt(name:str, text:list):
    
    if ".txt" in name:
        with open(docname, "w") as f:
            for line in text:
                f.write(line)
            f.close()

        os.startfile(docname)
    else:   
        docname = name + ".txt"
        with open(docname, "w") as f:
            for line in text:
                f.write(line)
            f.close()

    os.startfile(docname)


#Ejecucion
def main():
    
    interface()
    print("\n1. Transformar .docx a .txt")
    print("2. Crear un .txt en formato Aiken desde 0")
    
    while True:
        
        try:
            election = int(input("Elige una opción: "))

        except ValueError:
            election = int(input("Tipo de dato incorrecto. Elige una opción entre 1 y 2: "))

        except:
            election = int(input("Error inesperado, vuelve a intentarlo. Elige una opción entre 1 y 2: "))

        if election == 1 or election == 2:
            break

    if election == 1:

        while True:
            try:
                os.system("cls")
                interface1()
                documento = input("Introduce el nombre del documento(sin extensión): ")
                path = os.getcwd() + "\\" + documento + ".docx"
                addAnswer(path)
                text = getText(path)
                writeText(documento + ".txt", text)
                deleteFirstLine(documento + ".txt")
                deleteNumber(documento + ".txt")
                mayusc(documento + ".txt")
                print("Documento creado con éxito")
                os.startfile(documento + ".txt")
                exit()
            except docx.opc.exceptions.PackageNotFoundError:
                print("El documento no existe, vuelve a intentarlo.")

    elif election == 2:
        os.system("cls")
        interface2()
        name = str(input("\nIntroduce el nombre del archivo a crear: "))
        createTxt(name ,createQuestions())


main()
